#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Pandoc-based Markdown → DOCX converter.

Wraps pandoc subprocess, pre-processes Mermaid blocks via
:class:`~docx_pipeline.renderers.MermaidRenderer`, and post-processes the
generated ``.docx`` with python-docx to apply Chinese-friendly styling
(fonts, table borders, TOC field, page layout).

Design references
-----------------
* ``style_v16_docx.py`` — run-level + style-level font application, table-grid
  border injection, Word TOC field insertion.
* ``embed_mermaid_png.py`` — triple-check mmdc output validation pattern.
* ``PurePythonConverter`` — shared page-layout, core-properties, and TOC logic
  (deliberately duplicated here; see Appendix of Phase 2 design document).
"""

from __future__ import annotations

import logging
import os
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import List, Optional, Tuple

from docx import Document as DocxDocument
from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from docx_pipeline.config.schema import DocxPipelineConfig
from docx_pipeline.converters.base import AbstractConverter
from docx_pipeline.converters.shared import (
    _DEFAULT_HEADING_RGB,
    _DEFAULT_HEADING_SIZES,
    _PAGE_SIZES_CM,
    hex_to_rgb,
    set_cell_shading,
)

logger = logging.getLogger(__name__)

# Heading style names that pandoc produces (English locale)
_PANDOC_HEADING_STYLES = {
    "Heading 1", "Heading 2", "Heading 3",
    "Heading 4", "Heading 5", "Heading 6",
    "heading 1", "heading 2", "heading 3",
    "heading 4", "heading 5", "heading 6",
}


# ---------------------------------------------------------------------------
# PandocConverter
# ---------------------------------------------------------------------------


class PandocConverter(AbstractConverter):
    """Convert Markdown to DOCX using pandoc, with Mermaid pre-rendering and
    python-docx post-processing.

    Parameters
    ----------
    config : DocxPipelineConfig
        Fully-resolved pipeline configuration.
    extra_args : list of str, optional
        Additional pandoc CLI arguments (from ``--pandoc-args``).  Appended
        after ``config.pandoc.extra_args``.
    """

    def __init__(
        self,
        config: DocxPipelineConfig,
        extra_args: Optional[List[str]] = None,
    ) -> None:
        super().__init__(config)
        self.extra_args: List[str] = extra_args or []
        self.last_command: List[str] = []  # populated by _run_pandoc
        self._tempdir: Optional[tempfile.TemporaryDirectory] = None

    # ------------------------------------------------------------------
    # convert (AbstractContract)
    # ------------------------------------------------------------------

    def convert(self) -> Document:
        """Run the full pandoc conversion pipeline.

        Phases
        ------
        A. Setup work directory
        B. Pre-process Markdown (Mermaid → images)
        C. Invoke pandoc subprocess
        D. Load generated DOCX into python-docx
        E. Post-process (fonts, tables, TOC, footer)
        F. Set core properties
        """
        # A. Work directory
        work_dir = self._setup_work_dir()

        # B. Pre-process Markdown
        preprocessed_md = self._preprocess_markdown(work_dir)

        # C. Pandoc
        pandoc_output = work_dir / "pandoc_output.docx"
        self._run_pandoc(preprocessed_md, pandoc_output)

        # D. Load Document
        try:
            doc = DocxDocument(str(pandoc_output))
        except ValueError as exc:
            raise RuntimeError(
                f"pandoc 生成的 DOCX 无法打开：{exc}\n"
                "可能是 pandoc 版本或输入问题，请检查 preprocessed.md"
            ) from exc

        # E. Post-process
        self._postprocess(doc)

        # F. Core properties
        self._set_core_properties(doc)

        return doc

    # ------------------------------------------------------------------
    # Phase A: work directory
    # ------------------------------------------------------------------

    def _setup_work_dir(self) -> Path:
        """Create and return a unique working directory for intermediate files.

        Uses a UUID-named subdirectory under ``config.paths.work_dir`` (or a
        temp dir) so that concurrent conversions never share filenames.
        """
        import uuid
        wd = self.config.paths.work_dir
        if wd and wd.strip():
            p = (Path(wd) / f"run_{uuid.uuid4().hex[:8]}").resolve()
        else:
            p = Path(tempfile.gettempdir()) / f"docx_pipeline_{uuid.uuid4().hex[:8]}"
        p.mkdir(parents=True, exist_ok=True)
        self._tempdir = None  # we'll clean up manually
        return p

    # ------------------------------------------------------------------
    # Phase B: Markdown pre-processing
    # ------------------------------------------------------------------

    def _preprocess_markdown(self, work_dir: Path) -> Path:
        """Read the source Markdown, optionally render Mermaid blocks, and
        write the result to ``work_dir/preprocessed.md``.

        Returns the path to the preprocessed file.
        """
        md_source = Path(self.config.paths.md_source)
        if not md_source.is_absolute():
            md_source = (
                Path(self.config.project.root) / md_source
            ).resolve()
        md_source = md_source.resolve()

        if not md_source.exists():
            raise FileNotFoundError(
                f"Markdown source not found: {md_source}"
            )

        content = md_source.read_text(encoding="utf-8")

        # Mermaid pre-rendering (use same work_dir as pandoc for image resolution)
        if self.config.mermaid.enabled:
            try:
                from docx_pipeline.renderers import MermaidRenderer

                renderer = MermaidRenderer(self.config)
                content = renderer.render(content, work_dir=work_dir)
            except ImportError as exc:
                logger.warning(
                    "MermaidRenderer not available, skipping: %s", exc
                )
            except Exception as exc:
                logger.warning(
                    "Mermaid pre-rendering failed, keeping raw blocks: %s",
                    exc,
                )

        output_path = work_dir / "preprocessed.md"
        output_path.write_text(content, encoding="utf-8")
        return output_path

    # ------------------------------------------------------------------
    # Phase C: pandoc invocation
    # ------------------------------------------------------------------

    def _build_pandoc_args(self, md_path: Path) -> List[str]:
        """Construct the pandoc argument list from config.

        Parameters
        ----------
        md_path : Path
            The preprocessed markdown file path.  Used to derive
            ``--resource-path`` so that relative image references in the
            original markdown still resolve.
        """
        args = [
            "pandoc",
            "--from",
            (
                "markdown+pipe_tables+grid_tables+fenced_code_blocks"
                "+backtick_code_blocks+yaml_metadata_block+raw_html"
                "+superscript+subscript+strikeout+footnotes"
                "+definition_lists+example_lists+task_lists"
                "+multiline_tables+simple_tables"
            ),
            "--to",
            "docx",
            "--standalone",
            "--wrap",
            "preserve",
        ]

        # --resource-path: include the source dir, project root, and work dir
        # so that relative image paths in the original markdown still resolve
        # after the file is copied to work_dir/preprocessed.md.
        resource_dirs = []
        src_dir = str(Path(self.config.paths.md_source).parent.resolve())
        resource_dirs.append(src_dir)
        proj_root = str(Path(self.config.project.root).resolve())
        if proj_root not in resource_dirs:
            resource_dirs.append(proj_root)
        work_dir = str(md_path.parent.resolve())
        if work_dir not in resource_dirs:
            resource_dirs.append(work_dir)
        args += ["--resource-path", os.pathsep.join(resource_dirs)]

        # Reference docx for style inheritance
        ref = self.config.pandoc.reference_docx or self.config.paths.reference_docx
        if ref:
            args += ["--reference-doc", str(Path(ref).resolve())]

        # Metadata
        if self.config.project.name:
            args += ["--metadata", f"title={self.config.project.name}"]
        if self.config.version.date:
            args += ["--metadata", f"date={self.config.version.date}"]

        # Extra args: config first, CLI override last
        args.extend(self.config.pandoc.extra_args)
        args.extend(self.extra_args)

        return args

    def _run_pandoc(self, md_path: Path, output_path: Path) -> None:
        """Execute pandoc as a subprocess.

        Raises
        ------
        FileNotFoundError
            If pandoc is not installed.
        RuntimeError
            If pandoc exits non-zero.
        """
        # Pre-check
        if shutil.which("pandoc") is None:
            raise FileNotFoundError(
                "pandoc 未安装或不在 PATH 中。\n"
                "请通过 https://pandoc.org/installing.html 安装。"
            )

        cmd = self._build_pandoc_args(md_path) + [
            str(md_path), "-o", str(output_path)
        ]
        self.last_command = cmd

        logger.info("Running pandoc: %s", " ".join(cmd))

        env = os.environ.copy()
        env["PYTHONIOENCODING"] = "utf-8"

        try:
            proc = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=120,
                encoding="utf-8",
                errors="replace",
                env=env,
                cwd=str(md_path.parent),  # run pandoc from work_dir so relative image paths resolve
            )
            proc.check_returncode()
        except subprocess.TimeoutExpired:
            raise RuntimeError(
                "pandoc 超时（超过 120 秒）。"
                "文档可能过大，请尝试拆分或增加超时时间。"
            ) from None
        except subprocess.CalledProcessError as exc:
            stderr_tail = (
                exc.stderr.strip()[-500:]
                if exc.stderr
                else "(无 stderr)"
            )
            raise RuntimeError(
                f"pandoc 失败（退出码 {exc.returncode}）：{stderr_tail}"
            ) from exc

    # ------------------------------------------------------------------
    # Phase E: post-processing orchestrator
    # ------------------------------------------------------------------

    def _postprocess(self, doc: Document) -> None:
        """Apply all post-processing steps to the pandoc-generated document."""
        self._apply_page_layout(doc)
        self._apply_style_fonts(doc)
        self._normalize_block_text(doc)
        self._apply_paragraph_styles(doc)
        self._apply_run_fonts(doc)
        self._apply_code_block_shading(doc)
        self._apply_table_styles(doc)
        if self.config.styles.toc.enabled:
            self._apply_toc_field(doc)
        if self.config.version.label or self.config.version.date:
            self._apply_footer(doc)

    # ------------------------------------------------------------------
    # E1: Page layout
    # ------------------------------------------------------------------

    def _apply_page_layout(self, doc: Document) -> None:
        """Set page size, orientation, and margins from config."""
        section = doc.sections[0]

        size_key = self.config.page.size
        width_cm, height_cm = _PAGE_SIZES_CM.get(size_key, _PAGE_SIZES_CM["A4"])
        section.page_width = Cm(width_cm)
        section.page_height = Cm(height_cm)

        if self.config.page.orientation == "landscape":
            section.page_width, section.page_height = (
                section.page_height,
                section.page_width,
            )

        m = self.config.page.margins
        section.top_margin = Cm(m.top)
        section.bottom_margin = Cm(m.bottom)
        section.left_margin = Cm(m.left)
        section.right_margin = Cm(m.right)

    # ------------------------------------------------------------------
    # E2: Style-level fonts
    # ------------------------------------------------------------------

    def _apply_style_fonts(self, doc: Document) -> None:
        """Set font names on the *style definitions* for Normal and Heading
        1–6 so that new paragraphs inherit them."""
        latin = self.config.fonts.latin
        east_asian = self.config.fonts.east_asian

        for style_name in ["Normal"] + [
            f"Heading {n}" for n in range(1, 7)
        ]:
            style = doc.styles[style_name]
            rPr = style.element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.insert(0, rFonts)
            rFonts.set(qn("w:ascii"), latin)
            rFonts.set(qn("w:hAnsi"), latin)
            rFonts.set(qn("w:eastAsia"), east_asian)

    # ------------------------------------------------------------------
    # E2.1: Normalize Block Text → Normal (pandoc metadata styling fix)
    # ------------------------------------------------------------------

    def _normalize_block_text(self, doc: Document) -> None:
        """Convert pandoc ``Block Text`` paragraphs to ``Normal`` style.

        Pandoc uses ``Block Text`` for YAML metadata blocks, which often
        carries built-in left/right indentation in the default template.
        Converting to ``Normal`` produces a cleaner first page that matches
        user expectations (metadata flows as standard body text).
        """
        normalized = 0
        for para in doc.paragraphs:
            if para.style and para.style.name == "Block Text":
                para.style = doc.styles["Normal"]
                normalized += 1
        if normalized:
            logger.info("Normalized %d Block Text paragraph(s) → Normal", normalized)

    # ------------------------------------------------------------------
    # E2.5: Paragraph-level spacing
    # ------------------------------------------------------------------

    def _apply_paragraph_styles(self, doc: Document) -> None:
        """Apply line_spacing, space_after, and first_line_indent from
        ``config.styles.paragraph`` to Normal-style paragraphs.

        Also applies ``keep_with_next`` to heading paragraphs to prevent
        orphan headings, but **removes** it from headings immediately
        preceding an image — when the image is tall, the constraint forces
        both heading and image to the next page, creating a mostly-blank
        page (see memory: ``docx_keep_with_next_backfire``).
        """
        pc = self.config.styles.paragraph

        # Detect which paragraphs contain images (w:drawing elements)
        image_para_indices: set = set()
        for idx, para in enumerate(doc.paragraphs):
            drawings = para._element.findall(
                ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
            )
            if drawings:
                image_para_indices.add(idx)

        for idx, para in enumerate(doc.paragraphs):
            style_name = (para.style.name if para.style else "")

            # Apply "keep with next" to headings, UNLESS an image
            # paragraph follows within 1–3 paragraphs (blank lines
            # between heading and image in markdown → empty paragraphs
            # in DOCX).  Avoids the blank-page backfire effect.
            if "Heading" in style_name:
                pf = para.paragraph_format
                # Extended range (1..5): blank lines in markdown can produce
                # more than 3 empty paragraphs before an image.
                ahead = {idx + d for d in range(1, 6)}
                if ahead & image_para_indices:
                    pf.keep_with_next = False
                else:
                    pf.keep_with_next = True
                continue

            # Only apply spacing to body-text paragraphs, not code/TOC
            if "toc" in style_name.lower():
                continue
            if style_name in {"Source Code", "source code"}:
                continue

            pf = para.paragraph_format
            pf.line_spacing = pc.line_spacing
            pf.space_after = Pt(pc.space_after)
            if pc.first_line_indent:
                pf.first_line_indent = Cm(pc.first_line_indent)

    # ------------------------------------------------------------------
    # E3: Run-level fonts
    # ------------------------------------------------------------------

    def _apply_run_fonts(self, doc: Document) -> None:
        """Walk all paragraph runs and table-cell runs, applying font name,
        size, colour, and bold from config.

        Recognises pandoc-generated style names: ``Heading N``, ``Source Code``
        (code blocks), ``Hyperlink`` (links).
        """
        # Resolve heading overrides
        heading_overrides = {}
        for h_key, overrides in self.config.styles.heading.levels.items():
            try:
                level = int(h_key.lstrip("h"))
            except (ValueError, AttributeError):
                continue
            heading_overrides[level] = overrides

        # Resolve colours
        body_color = None
        if self.config.font_colors.body and self.config.font_colors.body != "auto":
            body_color = hex_to_rgb(self.config.font_colors.body)
        code_color = None
        if self.config.font_colors.code and self.config.font_colors.code != "auto":
            code_color = hex_to_rgb(self.config.font_colors.code)
        link_color = None
        if self.config.font_colors.link and self.config.font_colors.link != "auto":
            link_color = hex_to_rgb(self.config.font_colors.link)

        code_style_names = {"Source Code", "source code", "Verbatim Char"}

        # Walk paragraphs
        for para in doc.paragraphs:
            style_name = (para.style.name if para.style else "")
            heading_level = self._heading_level_from_style(style_name)
            is_code = style_name in code_style_names

            for run in para.runs:
                if heading_level is not None:
                    self._set_run_font_for_heading(
                        run, heading_level, heading_overrides
                    )
                elif is_code:
                    self._set_run_font(
                        run,
                        size=self.config.font_sizes.code,
                        color=code_color,
                        mono=True,
                    )
                elif "toc" in style_name.lower():
                    self._set_run_font(
                        run,
                        size=self.config.font_sizes.body,
                    )
                else:
                    self._set_run_font(
                        run,
                        size=self.config.font_sizes.body,
                        color=body_color,
                    )

            # Apply hyperlink colour to link-type runs in the paragraph
            for run in para.runs:
                rPr = run._element.find(qn("w:rPr"))
                if rPr is not None and rPr.find(qn("w:rStyle")) is not None:
                    rs = rPr.find(qn("w:rStyle"))
                    if rs is not None and "Hyperlink" in (rs.get(qn("w:val")) or ""):
                        if link_color:
                            run.font.color.rgb = RGBColor(*link_color)

        # Walk table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            self._set_run_font(
                                run,
                                size=self.config.font_sizes.table,
                            )

    def _heading_level_from_style(self, style_name: str) -> Optional[int]:
        """Extract heading level (1–6) from a style name, or None."""
        for level in range(1, 7):
            if f"Heading {level}" in style_name or f"heading {level}" in style_name:
                return level
        return None

    def _set_run_font_for_heading(
        self,
        run,
        level: int,
        overrides: dict,
    ) -> None:
        """Apply heading-specific font settings to a run."""
        lvl_overrides = overrides.get(level, {})

        # Font name
        ea = lvl_overrides.get("font_east_asian", self.config.fonts.east_asian)
        latin_f = lvl_overrides.get("font_latin", self.config.fonts.latin)

        # Font size
        heading_dict = self.config.font_sizes.headings
        key = f"h{level}"
        if "size" in lvl_overrides:
            size_pt = float(lvl_overrides["size"])
        elif key in heading_dict:
            size_pt = float(heading_dict[key])
        else:
            size_pt = _DEFAULT_HEADING_SIZES.get(level, 11.5)

        # Colour
        raw_color = lvl_overrides.get("color") or self.config.font_colors.heading
        color_rgb = None
        if raw_color and raw_color != "auto":
            color_rgb = hex_to_rgb(raw_color)
        else:
            color_rgb = _DEFAULT_HEADING_RGB

        # Bold
        bold = lvl_overrides.get("bold", True)

        self._set_run_font(
            run,
            size=size_pt,
            bold=bold,
            color=color_rgb,
            east_asian=ea,
            latin=latin_f,
        )

    # ------------------------------------------------------------------
    # Font helper
    # ------------------------------------------------------------------

    def _set_run_font(
        self,
        run,
        size: float = 10.5,
        bold: Optional[bool] = None,
        italic: Optional[bool] = None,
        color: Optional[Tuple[int, int, int]] = None,
        mono: bool = False,
        east_asian: Optional[str] = None,
        latin: Optional[str] = None,
    ) -> None:
        """Configure a run's font properties.

        Parameters
        ----------
        run : docx.text.run.Run
        size : float
            Font size in points.
        bold, italic : bool or None
            If *None* (default), preserve existing formatting.  Pass
            ``True`` or ``False`` to explicitly set.
        color : tuple of (R, G, B) or None
        mono : bool
            If True, use Consolas for Latin script.
        east_asian, latin : str or None
            Override font names; if None, read from config.
        """
        ea_font = east_asian or self.config.fonts.east_asian
        latin_font = latin or self.config.fonts.latin

        font = run.font
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)

        if mono:
            font.name = "Consolas"
            rFonts.set(qn("w:ascii"), "Consolas")
            rFonts.set(qn("w:hAnsi"), "Consolas")
            rFonts.set(qn("w:eastAsia"), ea_font)
        else:
            font.name = latin_font
            rFonts.set(qn("w:ascii"), latin_font)
            rFonts.set(qn("w:hAnsi"), latin_font)
            rFonts.set(qn("w:eastAsia"), ea_font)

        font.size = Pt(size)
        if bold is not None:
            font.bold = bold
        if italic is not None:
            font.italic = italic

        if color:
            font.color.rgb = RGBColor(*color)

    # ------------------------------------------------------------------
    # E3.5: Code block shading (v1.6.2 Deferred #3)
    # ------------------------------------------------------------------

    def _apply_code_block_shading(self, doc: Document) -> None:
        """Apply background shading to Source Code paragraphs.

        Adds ``w:shd`` to paragraphs whose style is ``Source Code``,
        using ``config.font_colors.code_block_bg`` (default ``#F5F5F5``).
        Table cells and TOC entries are skipped.
        """
        bg_color = self.config.font_colors.code_block_bg
        if not bg_color or bg_color == "auto":
            return

        color_hex = bg_color.lstrip("#")
        code_style_names = {"Source Code", "source code"}
        shaded = 0

        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ""
            if style_name not in code_style_names:
                continue

            # Add w:shd to the paragraph properties
            pPr = para._element.get_or_add_pPr()
            # Remove existing shading if any
            for child in list(pPr):
                if child.tag == qn("w:shd"):
                    pPr.remove(child)
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), color_hex)
            pPr.append(shd)
            shaded += 1

        if shaded:
            logger.info(
                "Applied code block shading (%s) to %d paragraph(s)",
                bg_color,
                shaded,
            )

    # ------------------------------------------------------------------
    # E4: Table styles
    # ------------------------------------------------------------------

    def _apply_table_styles(self, doc: Document) -> None:
        """Apply table style and explicit borders to all tables.

        Uses ``config.styles.table`` for style name, header shading, and
        bold/autofit preferences.  Pattern transplanted from
        ``style_v16_docx.py`` lines 33–59.
        """
        table_cfg = self.config.styles.table
        table_count = 0
        for table in doc.tables:
            tbl = table._tbl  # lxml <w:tbl> element

            # Find or create <w:tblPr>
            tblPr = tbl.find(qn("w:tblPr"))
            if tblPr is None:
                tblPr = OxmlElement("w:tblPr")
                tbl.insert(0, tblPr)

            # Remove existing tblStyle and inject configured style
            for child in list(tblPr):
                if child.tag == qn("w:tblStyle"):
                    tblPr.remove(child)
            tblStyle = OxmlElement("w:tblStyle")
            tblStyle.set(qn("w:val"), table_cfg.style)
            tblPr.insert(0, tblStyle)

            # Remove existing borders
            for child in list(tblPr):
                if child.tag == qn("w:tblBorders"):
                    tblPr.remove(child)

            # Add explicit borders
            borders = OxmlElement("w:tblBorders")
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                border = OxmlElement(f"w:{edge}")
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "4")
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), "auto")
                borders.append(border)
            tblPr.append(borders)

            # Remove existing width/layout to avoid duplicates
            for child in list(tblPr):
                if child.tag in (qn("w:tblW"), qn("w:tblLayout")):
                    tblPr.remove(child)

            # Set table width to 100% and enable autofit
            if table_cfg.autofit:
                tblW = OxmlElement("w:tblW")
                tblW.set(qn("w:w"), "5000")
                tblW.set(qn("w:type"), "pct")
                tblPr.append(tblW)
                tblLayout = OxmlElement("w:tblLayout")
                tblLayout.set(qn("w:type"), "autofit")
                tblPr.append(tblLayout)

            # Apply header-row shading from config
            if table_cfg.header_shading and table.rows:
                header_color = table_cfg.header_shading.lstrip("#")
                for cell in table.rows[0].cells:
                    set_cell_shading(cell, header_color)

            table_count += 1

        logger.info("Applied table styles to %d table(s)", table_count)

    # ------------------------------------------------------------------
    # E5: TOC field
    # ------------------------------------------------------------------

    def _apply_toc_field(self, doc: Document) -> None:
        """Locate the "目录" heading and insert a native Word TOC field,
        removing any hand-written TOC entries between the heading and the
        next heading.

        Pattern transplanted from ``style_v16_docx.py`` lines 61–134.
        """
        toc_title = self.config.styles.toc.title
        depth = self.config.styles.toc.depth

        toc_para_idx: Optional[int] = None
        toc_para_element = None

        # Find the TOC heading paragraph.
        # pandoc --number-sections prepends "N.M\t" to heading text, so we
        # match when the heading text *ends with* toc_title as well as exact
        # match (for unnumbered documents and the "目录" edge case where
        # pandoc renders the heading body as "1.1\t目录").
        for idx, para in enumerate(doc.paragraphs):
            style_name = para.style.name if para.style else ""
            if "Heading" not in style_name:
                continue
            text = para.text.strip()
            if text == toc_title or text.endswith("\t" + toc_title):
                toc_para_idx = idx
                toc_para_element = para._element
                break

        if toc_para_idx is None:
            logger.info(
                "TOC heading '%s' not found — skipping TOC field insertion",
                toc_title,
            )
            return

        # Remove hand-written TOC entries between the TOC heading and the
        # next heading.  Only delete paragraphs that look like TOC entries
        # (bullet/indented lines with links); leave body text untouched.
        import re as _re
        _TOC_LINE_RE = _re.compile(r'^\s*[-*+]\s+\[.+\]\(.+\)|^\s*\d+\.\s+\[.+\]\(.+\)')
        to_remove = []
        for idx in range(toc_para_idx + 1, len(doc.paragraphs)):
            para = doc.paragraphs[idx]
            style_name = para.style.name if para.style else ""
            if "Heading" in style_name:
                break
            text = para.text.strip()
            if not text:
                to_remove.append(para._element)  # blank separator lines
                continue
            if _TOC_LINE_RE.match(text):
                to_remove.append(para._element)
            # else: keep — looks like real body text, not TOC

        for elem in to_remove:
            elem.getparent().remove(elem)

        # Build the TOC field paragraph
        toc_p = OxmlElement("w:p")

        # fldChar begin
        run_begin = OxmlElement("w:r")
        fldChar_begin = OxmlElement("w:fldChar")
        fldChar_begin.set(qn("w:fldCharType"), "begin")
        run_begin.append(fldChar_begin)
        toc_p.append(run_begin)

        # instrText
        run_instr = OxmlElement("w:r")
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = f' TOC \\o "1-{depth}" \\h \\z \\u '
        run_instr.append(instrText)
        toc_p.append(run_instr)

        # fldChar separate
        run_sep = OxmlElement("w:r")
        fldChar_sep = OxmlElement("w:fldChar")
        fldChar_sep.set(qn("w:fldCharType"), "separate")
        run_sep.append(fldChar_sep)
        toc_p.append(run_sep)

        # Placeholder text
        run_placeholder = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:eastAsia"), self.config.fonts.east_asian)
        rPr.append(rFonts)
        run_placeholder.append(rPr)
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = "（在 Word 中右键此处 → 更新域，即可自动生成带页码的目录）"
        run_placeholder.append(t)
        toc_p.append(run_placeholder)

        # fldChar end
        run_end = OxmlElement("w:r")
        fldChar_end = OxmlElement("w:fldChar")
        fldChar_end.set(qn("w:fldCharType"), "end")
        run_end.append(fldChar_end)
        toc_p.append(run_end)

        # Insert after the TOC heading paragraph
        toc_para_element.addnext(toc_p)

        # Add a visible hint paragraph after the TOC field
        hint_p = OxmlElement("w:p")
        hint_r = OxmlElement("w:r")
        hint_rPr = OxmlElement("w:rPr")
        hint_color = OxmlElement("w:color")
        hint_color.set(qn("w:val"), "888888")
        hint_rPr.append(hint_color)
        hint_rFonts = OxmlElement("w:rFonts")
        hint_rFonts.set(qn("w:eastAsia"), self.config.fonts.east_asian)
        hint_rPr.append(hint_rFonts)
        hint_r.append(hint_rPr)
        hint_t = OxmlElement("w:t")
        hint_t.set(qn("xml:space"), "preserve")
        hint_t.text = (
            "（提示：打开文档后请右键目录 → 更新域，或按 F9，以显示页码。"
            "本行在更新后可删除。）"
        )
        hint_r.append(hint_t)
        hint_p.append(hint_r)
        toc_p.addnext(hint_p)

        logger.info("TOC field inserted after '%s' heading", toc_title)

    # ------------------------------------------------------------------
    # E6: Footer
    # ------------------------------------------------------------------

    def _apply_footer(self, doc: Document) -> None:
        """Add a version-status footer to the last section when version
        metadata is present."""
        label = self.config.version.label
        date = self.config.version.date

        if not label and not date:
            return

        if label and date:
            footer_text = f"本文档状态：{label}（{date}）"
        elif label:
            footer_text = f"本文档状态：{label}"
        else:
            footer_text = f"本文档日期：{date}"

        try:
            section = doc.sections[-1]
            footer = section.footer
            if footer.paragraphs:
                footer.paragraphs[0].clear()
                run = footer.paragraphs[0].add_run(footer_text)
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        except Exception:
            logger.debug("Footer insertion skipped (section/footer not available)")

    # ------------------------------------------------------------------
    # Phase F: core properties
    # ------------------------------------------------------------------

    def _set_core_properties(self, doc: Document) -> None:
        """Write document metadata into core properties."""
        if self.config.project.name:
            doc.core_properties.title = self.config.project.name
        if self.config.version.number:
            doc.core_properties.version = self.config.version.number
        if self.config.version.date:
            try:
                from datetime import datetime
                dt = datetime.fromisoformat(self.config.version.date)
                doc.core_properties.modified = dt
            except (ValueError, TypeError):
                pass  # skip if date string can't be parsed

    # ------------------------------------------------------------------
    # Cleanup
    # ------------------------------------------------------------------

    def cleanup(self) -> None:
        """Remove the temporary directory if one was created."""
        if self._tempdir is not None:
            try:
                self._tempdir.cleanup()
            except Exception:
                pass
            self._tempdir = None
