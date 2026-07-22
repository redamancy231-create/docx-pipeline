#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Config-driven pure-Python Markdown→DOCX converter.

Reads Markdown via :class:`MarkdownParser`, walks the resulting block list,
and assembles a ``python-docx`` :class:`Document` using every relevant knob
from :class:`~docx_pipeline.config.schema.DocxPipelineConfig`.

Design reference: ``md_to_docx_framework.py`` (V2.4, 366 lines) from the
AI协作框架 project — this module re-implements the same semantic logic
but is fully config-driven rather than hard-coded.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from docx import Document as DocxDocument
from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from docx_pipeline.config.schema import DocxPipelineConfig
from docx_pipeline.converters.base import AbstractConverter
from docx_pipeline.converters.shared import (
    _DEFAULT_HEADING_RGB,
    _DEFAULT_HEADING_SIZES,
    _PAGE_SIZES_CM,
    hex_to_rgb,
    set_cell_shading,
)
from docx_pipeline.converters.markdown_parser import (
    Block,
    BlockquoteBlock,
    Bold,
    CodeBlock,
    EmptyBlock,
    Format,
    HeadingBlock,
    HorizontalRuleBlock,
    InlineCode,
    ListBlock,
    MarkdownParser,
    ParagraphBlock,
    PlainText,
    TableBlock,
    parse_inline_formats,
)

# ═══════════════════════════════════════════════════════════════════════
# PurePythonConverter
# ═══════════════════════════════════════════════════════════════════════


class PurePythonConverter(AbstractConverter):
    """Convert a single Markdown file to DOCX using only Python logic.

    All styling decisions (fonts, sizes, colors, margins, page geometry)
    are read from the bound :class:`DocxPipelineConfig`.
    """

    _MATH_PLACEHOLDER_RE = re.compile(
        r"__DOCX_PIPELINE_MATH_\d{6}__"
    )
    _DISPLAY_MATH_RE = re.compile(
        r"(?<!\\)\$\$(.+?)(?<!\\)\$\$", re.DOTALL
    )
    _INLINE_MATH_RE = re.compile(
        r"(?<![\\$])\$(?!\$)([^$\n]+?)(?<!\\)\$(?!\$)"
    )
    _INLINE_CODE_SPAN_RE = re.compile(r"(`[^`\n]*`)")
    _FENCE_START_RE = re.compile(r"^\s*(`{3,}|~{3,})")

    # ------------------------------------------------------------------
    # convert
    # ------------------------------------------------------------------

    def convert(self) -> Document:
        """Run the full conversion pipeline.

        Raises
        ------
        FileNotFoundError
            If ``config.paths.md_source`` does not exist.
        ValueError
            If the Markdown content cannot be parsed (e.g. unclosed fence).
        """
        # 1. Resolve and read Markdown source
        md_path = Path(self.config.paths.md_source)
        if not md_path.is_absolute():
            md_path = (
                Path(self.config.project.root) / md_path
            ).resolve()
        md_path = md_path.resolve()

        if not md_path.exists():
            raise FileNotFoundError(
                f"Markdown source not found: {md_path}"
            )
        if not md_path.is_file():
            raise FileNotFoundError(
                f"Markdown source is not a file: {md_path}"
            )

        with open(md_path, "r", encoding="utf-8") as fh:
            content = fh.read()

        if not content.strip():
            raise ValueError(
                f"Markdown source is empty: {md_path}"
            )

        # 1.5. Pre-render Mermaid blocks (shared with Pandoc backend)
        _mermaid_work_dir = None  # for image resolution
        if self.config.mermaid.enabled:
            import logging as _logging
            _logger = _logging.getLogger(__name__)
            try:
                from docx_pipeline.renderers import MermaidRenderer
                import uuid, tempfile as _tempfile
                wd = self.config.paths.work_dir
                if wd and wd.strip():
                    _mermaid_work_dir = (Path(wd) / f"run_{uuid.uuid4().hex[:8]}").resolve()
                else:
                    _mermaid_work_dir = (Path(_tempfile.gettempdir()) / f"docx_pipeline_{uuid.uuid4().hex[:8]}").resolve()
                _mermaid_work_dir.mkdir(parents=True, exist_ok=True)
                renderer = MermaidRenderer(self.config)
                content = renderer.render(content, work_dir=_mermaid_work_dir)
            except ImportError:
                _logger.warning("MermaidRenderer not available", exc_info=True)
            except Exception:
                _logger.warning("Mermaid pre-rendering failed, keeping raw blocks", exc_info=True)

        # 1.75. Protect math from the Markdown parser. Display math is
        # extracted first so its dollar pairs cannot be consumed as inline math.
        content, self._math_placeholders = self._extract_math_placeholders(content)

        # Save resolved paths for image resolution later
        self._resolved_md_dir = md_path.parent.resolve()
        self._mermaid_work_dir = _mermaid_work_dir

        # 2. Parse into blocks
        parser = MarkdownParser(content)
        try:
            blocks = parser.parse()
        except ValueError:
            raise
        except Exception as exc:
            raise ValueError(
                f"Failed to parse Markdown from {md_path}: {exc}"
            ) from exc

        # 3. Build Document
        doc = self._create_document()

        # 4. Set core properties
        self._set_core_properties(doc)

        # 5. Walk blocks and emit content
        self._build_content(doc, blocks)

        return doc

    # ------------------------------------------------------------------
    # Document scaffolding
    # ------------------------------------------------------------------

    def _create_document(self) -> Document:
        """Create and configure the ``Document`` object (page, margins)."""
        doc = DocxDocument()

        section = doc.sections[0]

        # Page size
        size_key = self.config.page.size
        width_cm, height_cm = _PAGE_SIZES_CM.get(
            size_key, _PAGE_SIZES_CM["A4"]
        )
        section.page_width = Cm(width_cm)
        section.page_height = Cm(height_cm)

        # Orientation swap
        if self.config.page.orientation == "landscape":
            section.page_width, section.page_height = (
                section.page_height,
                section.page_width,
            )

        # Margins
        m = self.config.page.margins
        section.top_margin = Cm(m.top)
        section.bottom_margin = Cm(m.bottom)
        section.left_margin = Cm(m.left)
        section.right_margin = Cm(m.right)

        return doc

    def _set_core_properties(self, doc: Document) -> None:
        """Write core document metadata."""
        doc.core_properties.title = self.config.project.name
        if self.config.version.number:
            doc.core_properties.version = self.config.version.number

    # ------------------------------------------------------------------
    # Math extraction and OMML generation
    # ------------------------------------------------------------------

    def _extract_math_placeholders(
        self, content: str
    ) -> Tuple[str, Dict[str, Tuple[str, bool]]]:
        """Replace LaTeX math spans with parser-safe placeholders.

        Display math is extracted before inline math. Fenced code blocks and
        inline code spans are deliberately left untouched so code examples do
        not become equations.
        """
        placeholders: Dict[str, Tuple[str, bool]] = {}
        counter = 0

        def new_placeholder(latex: str, is_display: bool) -> str:
            nonlocal counter
            counter += 1
            placeholder = f"__DOCX_PIPELINE_MATH_{counter:06d}__"
            placeholders[placeholder] = (latex, is_display)
            return placeholder

        def process_prose(prose: str) -> str:
            # Inline code is protected before either dollar-delimited form is
            # considered. The Markdown parser handles the backticks later.
            parts = self._INLINE_CODE_SPAN_RE.split(prose)
            for index in range(0, len(parts), 2):
                part = parts[index]

                def replace_display(match: re.Match[str]) -> str:
                    latex = match.group(1).strip()
                    if not latex:
                        return match.group(0)
                    placeholder = new_placeholder(latex, True)
                    # Blank lines force a display placeholder to be parsed as
                    # its own block even when the source delimiters were not.
                    return f"\n\n{placeholder}\n\n"

                part = self._DISPLAY_MATH_RE.sub(replace_display, part)

                def replace_inline(match: re.Match[str]) -> str:
                    latex = match.group(1).strip()
                    if not latex:
                        return match.group(0)
                    return new_placeholder(latex, False)

                parts[index] = self._INLINE_MATH_RE.sub(replace_inline, part)
            return "".join(parts)

        # Process prose chunks while copying fenced code blocks byte-for-byte.
        output: List[str] = []
        prose_lines: List[str] = []
        fence_char: Optional[str] = None
        fence_length = 0

        def flush_prose() -> None:
            if prose_lines:
                output.append(process_prose("".join(prose_lines)))
                prose_lines.clear()

        for line in content.splitlines(keepends=True):
            if fence_char is None:
                fence_match = self._FENCE_START_RE.match(line)
                if fence_match:
                    flush_prose()
                    marker_text = fence_match.group(1)
                    fence_char = marker_text[0]
                    fence_length = len(marker_text)
                    output.append(line)
                else:
                    prose_lines.append(line)
                continue

            output.append(line)
            stripped = line.strip()
            if (
                len(stripped) >= fence_length
                and stripped
                and set(stripped) == {fence_char}
            ):
                fence_char = None
                fence_length = 0

        flush_prose()
        return "".join(output), placeholders

    def _latex_to_omml(self, latex: str):
        """Return an ``m:oMath`` tree, or ``None`` when conversion fails."""
        from docx_pipeline.converters.mathml2omml import (
            MathConversionError,
            latex_to_omml,
        )

        try:
            return latex_to_omml(latex, self.config.font_sizes.body)
        except MathConversionError:
            return None
    @staticmethod
    def _make_math_run(text: str):
        run = OxmlElement("m:r")
        text_element = OxmlElement("m:t")
        if text.startswith(" ") or text.endswith(" "):
            text_element.set(qn("xml:space"), "preserve")
        text_element.text = text
        run.append(text_element)
        return run

    @staticmethod
    def _make_math_argument(tag: str, nodes: List[object]):
        argument = OxmlElement(tag)
        for node in nodes:
            argument.append(node)
        return argument

    def _add_display_math(self, doc: Document, latex: str) -> None:
        math = self._latex_to_omml(latex)
        if math is None:
            self._add_paragraph(doc, f"$${latex}$$", [])
            return

        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(
            self.config.styles.paragraph.space_after
        )
        math_paragraph = OxmlElement("m:oMathPara")
        properties = OxmlElement("m:oMathParaPr")
        justification = OxmlElement("m:jc")
        justification.set(qn("m:val"), "center")
        properties.append(justification)
        math_paragraph.append(properties)
        math_paragraph.append(math)
        paragraph._p.append(math_paragraph)

    def _append_text_with_math(
        self,
        paragraph,
        text: str,
        *,
        size: float,
        bold: Optional[bool] = None,
        italic: Optional[bool] = None,
        color: Optional[Tuple[int, int, int]] = None,
        mono: bool = False,
        level_override: Optional[dict] = None,
    ) -> None:
        """Append styled text and inline OMML in exact source order."""
        placeholders = getattr(self, "_math_placeholders", {})
        cursor = 0

        def add_text(value: str) -> None:
            if not value:
                return
            run = paragraph.add_run(value)
            self._set_run_font(
                run,
                size=size,
                bold=bold,
                italic=italic,
                color=color,
                mono=mono,
                level_override=level_override,
            )

        for match in self._MATH_PLACEHOLDER_RE.finditer(text):
            add_text(text[cursor:match.start()])
            placeholder = match.group(0)
            math_data = placeholders.get(placeholder)
            if math_data is None or math_data[1]:
                add_text(placeholder)
            else:
                latex = math_data[0]
                math = self._latex_to_omml(latex)
                if math is None:
                    add_text(f"${latex}$")
                else:
                    paragraph._p.append(math)
            cursor = match.end()
        add_text(text[cursor:])

    # ------------------------------------------------------------------
    # Block dispatcher
    # ------------------------------------------------------------------

    def _build_content(self, doc: Document, blocks: List[Block]) -> None:
        """Iterate *blocks* and delegate to the appropriate ``_add_*`` method."""
        toc_inserted = False

        # Pre-scan: check if heading text matches TOC trigger
        toc_title_texts = ["目录"]

        for block in blocks:
            # Check for TOC-trigger heading before dispatching
            if (
                self.config.styles.toc.enabled
                and not toc_inserted
                and isinstance(block, HeadingBlock)
            ):
                if block.text.strip() in toc_title_texts:
                    self._insert_toc_field(doc)
                    toc_inserted = True
                    continue   # skip the heading itself (TOC field replaces it)

            # Dispatch
            if isinstance(block, HeadingBlock):
                self._add_heading(doc, block.level, block.text)
            elif isinstance(block, ParagraphBlock):
                math_data = getattr(self, "_math_placeholders", {}).get(
                    block.text.strip()
                )
                if math_data is not None and math_data[1]:
                    self._add_display_math(doc, math_data[0])
                else:
                    self._add_paragraph_or_image(
                        doc, block.text, block.formats
                    )
            elif isinstance(block, CodeBlock):
                self._add_code_block(doc, block.lines, block.language)
            elif isinstance(block, TableBlock):
                self._add_table(doc, block.rows)
            elif isinstance(block, BlockquoteBlock):
                self._add_blockquote(doc, block.text)
            elif isinstance(block, ListBlock):
                self._add_list(doc, block.items, block.ordered, block.indent)
            elif isinstance(block, HorizontalRuleBlock):
                self._add_horizontal_rule(doc)
            elif isinstance(block, EmptyBlock):
                # Skip explicit empties at this level; spacing is managed
                # by paragraph-level space_before / space_after.
                pass

    # ------------------------------------------------------------------
    # Image-aware paragraph dispatch
    # ------------------------------------------------------------------

    _IMAGE_RE = re.compile(r"^!\[(.+?)\]\((.+?)\)$")

    def _add_paragraph_or_image(
        self, doc: Document, text: str, formats: List[Format]
    ) -> None:
        """Detect stand-alone ``![caption](path)`` image references and embed
        the actual image; fall back to normal paragraph otherwise."""
        m = self._IMAGE_RE.match(text.strip())
        if m:
            caption = m.group(1)
            img_path = m.group(2)
            # Try relative to md_source dir first, then relative to mermaid work dir
            md_dir = self._resolved_md_dir
            full_path = (md_dir / img_path).resolve()
            if not full_path.is_file() and self._mermaid_work_dir is not None:
                full_path = (self._mermaid_work_dir / img_path).resolve()
            if full_path.is_file():
                self._add_image(doc, str(full_path), caption)
                return
        # Not an image reference or file not found — normal paragraph
        self._add_paragraph(doc, text, formats)

    def _add_image(self, doc: Document, path: str, caption: str = "") -> None:
        """Insert an image into the document, auto-sizing to page width."""
        import logging as _logging
        _img_logger = _logging.getLogger(__name__)
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            inline_shape = run.add_picture(path)
            # Constrain width to usable page width
            self._constrain_image_width(doc, inline_shape)
            if caption:
                cap_p = doc.add_paragraph()
                cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_run = cap_p.add_run(caption)
                cap_run.font.size = Pt(self.config.font_sizes.body - 1)
                cap_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        except Exception:
            _img_logger.warning("Cannot embed image: %s", path, exc_info=True)
            self._add_paragraph(doc, f"[Image: {caption}]({path})", [])

    def _constrain_image_width(self, doc: Document, inline_shape) -> None:
        """Scale an image down to fit within the usable page width.  Does not
        enlarge small images.  Reads the actual section dimensions so that
        landscape orientation and custom page sizes are handled correctly."""
        try:
            section = doc.sections[0]
            # Usable width in EMU from the actual section (portrait, landscape,
            # or custom — already applied by _create_document)
            usable_emu = section.page_width - section.left_margin - section.right_margin
            if inline_shape.width > usable_emu:
                ratio = usable_emu / inline_shape.width
                inline_shape.width = usable_emu
                inline_shape.height = int(inline_shape.height * ratio)
        except Exception:
            pass

    # ------------------------------------------------------------------
    # Heading
    # ------------------------------------------------------------------

    def _add_heading(self, doc: Document, level: int, text: str) -> None:
        """Add a styled heading paragraph.

        Font size: ``config.font_sizes.headings["h<N>"]`` or built-in default.
        Colour: ``config.font_colors.heading`` or ``_DEFAULT_HEADING_RGB``.
        Font: ``config.fonts`` heading overrides from
              ``config.styles.heading.levels`` if present.
        """
        level = max(1, min(6, level))

        p = doc.add_heading(level=level)

        # Resolve size: per-level override > font_sizes.headings > default
        heading_dict = self.config.font_sizes.headings
        key = f"h{level}"
        level_overrides = self.config.styles.heading.levels.get(key, {})
        if "size" in level_overrides:
            size_pt = float(level_overrides["size"])
        elif key in heading_dict:
            size_pt = float(heading_dict[key])
        else:
            size_pt = _DEFAULT_HEADING_SIZES.get(level, 11.5)

        # Resolve colour: per-level override > font_colors.heading > default
        color_rgb = None
        raw_color = self.config.font_colors.heading
        if "color" in level_overrides:
            color_rgb = hex_to_rgb(level_overrides["color"])
        elif raw_color and raw_color != "auto":
            color_rgb = hex_to_rgb(raw_color)
        else:
            color_rgb = _DEFAULT_HEADING_RGB

        # Resolve font name: per-level override > config.fonts
        if "font_east_asian" in level_overrides:
            level_overrides = dict(level_overrides)  # don't mutate config
        bold = level_overrides.get("bold", True)
        italic = level_overrides.get("italic", False)

        self._append_text_with_math(
            p,
            text,
            size=size_pt,
            bold=bold,
            italic=italic,
            color=color_rgb,
            mono=False,
            level_override=level_overrides,
        )

        # Paragraph formatting
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h_cfg = self.config.styles.heading
        sb = h_cfg.space_before.get(key, 8.0)
        p.paragraph_format.space_before = Pt(sb)
        p.paragraph_format.space_after = Pt(h_cfg.space_after)

    # ------------------------------------------------------------------
    # Paragraph
    # ------------------------------------------------------------------

    def _add_paragraph(
        self,
        doc: Document,
        text: str,
        formats: List[Format],
        *,
        bold: bool = False,
        italic: bool = False,
        color: Optional[Tuple[int, int, int]] = None,
        size: Optional[float] = None,
        indent: bool = False,
        mono: bool = False,
        left_indent_cm: Optional[float] = None,
    ) -> None:
        """Add a paragraph, optionally with inline formatting runs.

        If *formats* is empty, a single run is created from *text* with the
        provided style kwargs.  Otherwise each ``Format`` token becomes a
        separate run.
        """
        p = doc.add_paragraph()

        # Paragraph-level formatting
        p.paragraph_format.line_spacing = self.config.styles.paragraph.line_spacing
        p.paragraph_format.space_after = Pt(
            self.config.styles.paragraph.space_after
        )
        if indent or self.config.styles.paragraph.first_line_indent:
            p.paragraph_format.first_line_indent = Cm(
                self.config.styles.paragraph.first_line_indent or 0.74
            )
        if left_indent_cm is not None:
            p.paragraph_format.left_indent = Cm(left_indent_cm)

        body_size = size or self.config.font_sizes.body

        # Apply body colour from config when not explicitly overridden
        if color is None and not mono:
            bc = self.config.font_colors.body
            if bc and bc != "auto":
                color = hex_to_rgb(bc)

        if not formats:
            self._append_text_with_math(
                p,
                text,
                size=body_size,
                bold=bold,
                italic=italic,
                color=color,
                mono=mono,
            )
        else:
            for fmt in formats:
                if isinstance(fmt, Bold):
                    self._append_text_with_math(
                        p,
                        fmt.text,
                        size=body_size,
                        bold=True,
                        italic=italic,
                        color=color,
                        mono=mono,
                    )
                elif isinstance(fmt, InlineCode):
                    self._append_text_with_math(
                        p,
                        fmt.text,
                        size=body_size - 0.5,
                        bold=False,
                        italic=italic,
                        color=hex_to_rgb(
                            self.config.font_colors.code
                        ) if self.config.font_colors.code != "auto" else None,
                        mono=True,
                    )
                elif isinstance(fmt, PlainText):
                    self._append_text_with_math(
                        p,
                        fmt.text,
                        size=body_size,
                        bold=bold,
                        italic=italic,
                        color=color,
                        mono=mono,
                    )

    # ------------------------------------------------------------------
    # Code block
    # ------------------------------------------------------------------

    def _add_code_block(
        self,
        doc: Document,
        lines: List[str],
        language: str,
    ) -> None:
        """Add a shaded, monospaced code block."""
        code_size = self.config.font_sizes.code
        code_color = None
        if self.config.font_colors.code != "auto":
            code_color = hex_to_rgb(self.config.font_colors.code)

        for line in lines:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

            # Light grey background via paragraph shading
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), self.config.font_colors.code_block_bg.lstrip("#"))
            shd.set(qn("w:val"), "clear")
            pPr.append(shd)

            run = p.add_run(line if line else " ")
            self._set_run_font(
                run, size=code_size, color=code_color, mono=True
            )

        # Spacer after code block
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(2)
        spacer.paragraph_format.space_after = Pt(2)

    # ------------------------------------------------------------------
    # Table
    # ------------------------------------------------------------------

    def _add_table(self, doc: Document, rows: List[List[str]]) -> None:
        """Add a styled table with header-row shading."""
        if not rows:
            return

        num_cols = max(len(r) for r in rows)
        table = doc.add_table(rows=len(rows), cols=num_cols)
        table.style = self.config.styles.table.style
        table.autofit = self.config.styles.table.autofit
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        table_size = self.config.font_sizes.table
        header_shading = self.config.styles.table.header_shading.lstrip("#")

        for i, row in enumerate(rows):
            for j in range(num_cols):
                cell = table.rows[i].cells[j]
                cell.text = ""
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_before = Pt(1)
                paragraph.paragraph_format.space_after = Pt(1)
                if j < len(row):
                    is_header = (
                        i == 0 and self.config.styles.table.header_bold
                    )
                    self._append_text_with_math(
                        paragraph,
                        row[j],
                        size=table_size,
                        bold=is_header,
                    )
                if i == 0:
                    set_cell_shading(cell, header_shading)

        # Spacer after table
        doc.add_paragraph()

    # ------------------------------------------------------------------
    # Blockquote
    # ------------------------------------------------------------------

    def _add_blockquote(
        self,
        doc: Document,
        text: str,
    ) -> None:
        """Add an indented, italic, grey blockquote paragraph."""
        fmt_tokens = parse_inline_formats(text)

        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = self.config.styles.paragraph.line_spacing

        bq_color = hex_to_rgb(self.config.font_colors.blockquote)
        body_size = self.config.font_sizes.body

        if not fmt_tokens:
            self._append_text_with_math(
                p,
                text,
                size=body_size,
                italic=True,
                color=bq_color,
            )
        else:
            for fmt in fmt_tokens:
                if isinstance(fmt, Bold):
                    self._append_text_with_math(
                        p,
                        fmt.text,
                        size=body_size,
                        bold=True,
                        italic=True,
                        color=bq_color,
                    )
                elif isinstance(fmt, InlineCode):
                    self._append_text_with_math(
                        p,
                        fmt.text,
                        size=body_size - 0.5,
                        italic=False,
                        mono=True,
                        color=bq_color,
                    )
                elif isinstance(fmt, PlainText):
                    self._append_text_with_math(
                        p,
                        fmt.text,
                        size=body_size,
                        italic=True,
                        color=bq_color,
                    )

    # ------------------------------------------------------------------
    # List
    # ------------------------------------------------------------------

    def _add_list(
        self,
        doc: Document,
        items: List[str],
        ordered: bool,
        indent: int,
    ) -> None:
        """Add a contiguous list block using Word native list styles."""
        body_size = self.config.font_sizes.body
        list_style = "List Number" if ordered else "List Bullet"

        for item_text in items:
            fmt_tokens = parse_inline_formats(item_text)

            p = doc.add_paragraph(style=list_style)
            p.paragraph_format.left_indent = Cm(0.74 + indent * 0.5)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)

            if not fmt_tokens:
                self._append_text_with_math(
                    p, item_text, size=body_size
                )
            else:
                for fmt in fmt_tokens:
                    if isinstance(fmt, Bold):
                        self._append_text_with_math(
                            p, fmt.text, size=body_size, bold=True
                        )
                    elif isinstance(fmt, InlineCode):
                        self._append_text_with_math(
                            p, fmt.text, size=body_size - 0.5, mono=True
                        )
                    elif isinstance(fmt, PlainText):
                        self._append_text_with_math(
                            p, fmt.text, size=body_size
                        )

    # ------------------------------------------------------------------
    # Horizontal rule
    # ------------------------------------------------------------------

    def _add_horizontal_rule(self, doc: Document) -> None:
        """Insert a thin horizontal line via paragraph bottom border."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)

        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:color"), self.config.font_colors.horizontal_rule.lstrip("#"))
        pBdr.append(bottom)
        pPr.append(pBdr)

    # ------------------------------------------------------------------
    # TOC field
    # ------------------------------------------------------------------

    def _insert_toc_field(self, doc: Document) -> None:
        """Insert a native Word TOC field.

        The user must right-click → Update Field (F9) in Word to populate
        page numbers.  The TOC depth is taken from ``config.styles.toc.depth``.
        """
        depth = self.config.styles.toc.depth
        toc_title = self.config.styles.toc.title

        # Add a heading for the TOC so it appears in the navigation pane
        p_title = doc.add_paragraph()
        run_title = p_title.add_run(toc_title)
        self._set_run_font(
            run_title,
            size=self.config.font_sizes.body + 2,
            bold=True,
            color=_DEFAULT_HEADING_RGB,
        )
        p_title.paragraph_format.space_after = Pt(6)

        # Build the field
        toc_paragraph = doc.add_paragraph()
        toc_paragraph.style = doc.styles["Normal"]

        # fldChar begin
        run1 = toc_paragraph.add_run()
        fldChar_begin = OxmlElement("w:fldChar")
        fldChar_begin.set(qn("w:fldCharType"), "begin")
        run1._r.append(fldChar_begin)

        # instrText
        run2 = toc_paragraph.add_run()
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = f' TOC \\o "1-{depth}" \\h \\z \\u '
        run2._r.append(instrText)

        # fldChar separate
        run3 = toc_paragraph.add_run()
        fldChar_separate = OxmlElement("w:fldChar")
        fldChar_separate.set(qn("w:fldCharType"), "separate")
        run3._r.append(fldChar_separate)

        # Placeholder text
        run4 = toc_paragraph.add_run()
        run4_text = OxmlElement("w:t")
        run4_text.text = "（在 Word 中右键此处 → 更新域，即可自动生成带页码的目录）"
        run4._r.append(run4_text)

        # fldChar end
        run5 = toc_paragraph.add_run()
        fldChar_end = OxmlElement("w:fldChar")
        fldChar_end.set(qn("w:fldCharType"), "end")
        run5._r.append(fldChar_end)

        # Spacer
        doc.add_paragraph()

    # ------------------------------------------------------------------
    # Font helper
    # ------------------------------------------------------------------

    def _set_run_font(
        self,
        run,
        size: float = 10.5,
        bold: Optional[bool] = None,
        italic: Optional[bool] = None,
        color: Optional[Tuple[int, int, int]] = None,
        mono: bool = False,
        level_override: Optional[dict] = None,
    ) -> None:
        """Configure a run's font from config.

        Parameters
        ----------
        run : docx.text.run.Run
        size : float
            Font size in points.
        bold, italic : bool or None
            If *None* (default), preserve existing formatting.
        color : tuple of (R, G, B) or None
        mono : bool
            If True, use ``Consolas`` (Latin) while keeping
            ``config.fonts.east_asian`` for CJK.
        level_override : dict, optional
            Per-heading-level overrides from
            ``config.styles.heading.levels["h<N>"]``.
        """
        font = run.font
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)

        # Resolve font names (may be overridden per heading level)
        ea_font = self.config.fonts.east_asian
        latin_font = self.config.fonts.latin

        if level_override:
            ea_font = level_override.get("font_east_asian", ea_font)
            latin_font = level_override.get("font_latin", latin_font)

        if mono:
            font.name = "Consolas"
            rFonts.set(qn("w:ascii"), "Consolas")
            rFonts.set(qn("w:hAnsi"), "Consolas")
            rFonts.set(qn("w:eastAsia"), ea_font)
        else:
            font.name = latin_font
            rFonts.set(qn("w:ascii"), latin_font)
            rFonts.set(qn("w:hAnsi"), latin_font)
            rFonts.set(qn("w:eastAsia"), ea_font)

        font.size = Pt(size)
        if bold is not None:
            font.bold = bold
        if italic is not None:
            font.italic = italic

        if color:
            font.color.rgb = RGBColor(*color)
