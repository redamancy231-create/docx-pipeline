"""Minimal smoke tests for docx_pipeline."""

import os
import sys
import tempfile
from pathlib import Path

import pytest

# Ensure the package is importable
sys.path.insert(0, str(Path(__file__).resolve().parents[1]))


class TestConfig:
    """Configuration loading and validation."""

    def test_load_default_template(self):
        from docx_pipeline.config.defaults import get_template, TEMPLATES
        for name in ["default", "academic", "report", "strategy"]:
            data = get_template(name)
            assert isinstance(data, dict)
            assert "project" in data
            assert "paths" in data
            assert data["paths"]["md_source"]
            assert data["paths"]["docx_output"]

    def test_from_dict_roundtrip(self):
        from docx_pipeline.config.defaults import get_template
        from docx_pipeline.config.schema import DocxPipelineConfig
        data = get_template("default")
        cfg = DocxPipelineConfig.from_dict(data)
        assert cfg.project.name == "default_project"
        assert cfg.fonts.east_asian == "微软雅黑"
        # Round-trip
        exported = cfg.to_dict()
        assert exported["project"]["name"] == "default_project"

    def test_template_pandoc_defaults(self):
        """Verify #2: only report template enables pandoc/mermaid by default."""
        from docx_pipeline.config.defaults import get_template
        for name in ["default", "academic", "strategy"]:
            data = get_template(name)
            assert data["pandoc"]["enabled"] is False, f"{name} pandoc should be off"
            assert data["mermaid"]["enabled"] is False, f"{name} mermaid should be off"
        report = get_template("report")
        assert report["pandoc"]["enabled"] is True
        assert report["mermaid"]["enabled"] is True

    def test_config_rejects_unknown_version_type(self):
        """Verify #1: version as plain string should raise TypeError."""
        from docx_pipeline.config.schema import DocxPipelineConfig
        data = {
            "project": {"name": "test", "root": "."},
            "paths": {"md_source": "test.md", "docx_output": "out.docx"},
            "version": "1.0",  # old-style string, should fail
        }
        with pytest.raises(TypeError):
            DocxPipelineConfig.from_dict(data)


class TestMarkdownParser:
    """Markdown block parsing."""

    def test_heading_parsing(self):
        from docx_pipeline.converters.markdown_parser import MarkdownParser, HeadingBlock
        parser = MarkdownParser("# Hello\n## World\nplain text\n")
        blocks = parser.parse()
        headings = [b for b in blocks if isinstance(b, HeadingBlock)]
        assert len(headings) == 2
        assert headings[0].level == 1
        assert headings[0].text == "Hello"

    def test_code_fence(self):
        from docx_pipeline.converters.markdown_parser import MarkdownParser, CodeBlock
        md = "```python\nprint('hi')\n```\n"
        blocks = MarkdownParser(md).parse()
        codes = [b for b in blocks if isinstance(b, CodeBlock)]
        assert len(codes) == 1
        assert codes[0].language == "python"
        assert codes[0].lines == ["print('hi')"]

    def test_unclosed_fence_raises(self):
        from docx_pipeline.converters.markdown_parser import MarkdownParser
        with pytest.raises(ValueError, match="Unclosed code fence"):
            MarkdownParser("```python\nprint('hi')\n").parse()

    def test_table_parsing(self):
        from docx_pipeline.converters.markdown_parser import MarkdownParser, TableBlock
        md = "| A | B |\n|---|---|\n| 1 | 2 |\n"
        blocks = MarkdownParser(md).parse()
        tables = [b for b in blocks if isinstance(b, TableBlock)]
        assert len(tables) == 1
        assert tables[0].rows == [["A", "B"], ["1", "2"]]

    def test_yaml_frontmatter_skipped(self):
        from docx_pipeline.converters.markdown_parser import MarkdownParser, HeadingBlock
        md = "---\ntitle: Test\n---\n# Real Content\n"
        blocks = MarkdownParser(md).parse()
        headings = [b for b in blocks if isinstance(b, HeadingBlock)]
        assert len(headings) == 1
        assert headings[0].text == "Real Content"


class TestCLI:
    """CLI smoke tests (no external dependencies needed)."""

    def test_cli_help(self):
        from click.testing import CliRunner
        from docx_pipeline.cli import cli
        runner = CliRunner()
        result = runner.invoke(cli, ["--help"])
        assert result.exit_code == 0
        assert "init" in result.output
        assert "convert" in result.output
        assert "validate" in result.output

    def test_init_creates_config(self):
        from click.testing import CliRunner
        from docx_pipeline.cli import cli
        runner = CliRunner()
        with tempfile.TemporaryDirectory() as tmp:
            result = runner.invoke(cli, [
                "init", "-d", tmp, "-n", "test-project", "-t", "default",
            ])
            assert result.exit_code == 0
            assert os.path.exists(os.path.join(tmp, "project.yaml"))

    def test_validate_fails_on_missing_md(self):
        from click.testing import CliRunner
        from docx_pipeline.cli import cli
        runner = CliRunner()
        with tempfile.TemporaryDirectory() as tmp:
            # Create a valid project.yaml with non-existent md_source
            import yaml
            config = {
                "project": {"name": "test", "root": tmp},
                "paths": {
                    "md_source": os.path.join(tmp, "nonexistent.md"),
                    "docx_output": os.path.join(tmp, "out.docx"),
                },
                "pandoc": {"enabled": False, "extra_args": [], "reference_docx": ""},
                "mermaid": {"enabled": False, "image": {}, "render": {}},
                "version": {"number": "1.0.0", "label": "", "date": ""},
            }
            config_path = os.path.join(tmp, "project.yaml")
            with open(config_path, "w") as f:
                yaml.dump(config, f)
            result = runner.invoke(cli, ["validate", "-c", config_path])
            assert result.exit_code != 0  # should fail on missing md_source

    def test_convert_dry_run(self):
        from click.testing import CliRunner
        from docx_pipeline.cli import cli
        import yaml
        runner = CliRunner()
        with tempfile.TemporaryDirectory() as tmp:
            # Create a minimal markdown file
            md_path = os.path.join(tmp, "test.md")
            with open(md_path, "w") as f:
                f.write("# Hello\n\nWorld.\n")
            config = {
                "project": {"name": "test", "root": tmp},
                "paths": {
                    "md_source": md_path,
                    "docx_output": os.path.join(tmp, "out.docx"),
                },
                "pandoc": {"enabled": False, "extra_args": [], "reference_docx": ""},
                "mermaid": {"enabled": False, "image": {}, "render": {}},
                "version": {"number": "1.0.0", "label": "", "date": ""},
            }
            config_path = os.path.join(tmp, "project.yaml")
            with open(config_path, "w") as f:
                yaml.dump(config, f)
            result = runner.invoke(cli, [
                "convert", "-c", config_path, "--dry-run",
            ])
            assert result.exit_code == 0
            assert "DRY-RUN" in result.output


class TestBackup:
    """Backup rotation (#7)."""

    def test_backup_rotation(self):
        from docx_pipeline.config.defaults import get_template
        from docx_pipeline.config.schema import DocxPipelineConfig
        from docx_pipeline.converters.base import AbstractConverter

        data = get_template("default")
        data["backup"]["enabled"] = True
        data["backup"]["max_backups"] = 3
        cfg = DocxPipelineConfig.from_dict(data)

        class _TestConverter(AbstractConverter):
            def convert(self):
                from docx import Document
                return Document()

        with tempfile.TemporaryDirectory() as tmp:
            out = Path(tmp) / "out.docx"
            out.write_text("v1")
            converter = _TestConverter(cfg)
            converter._rotate_backups(out)
            bak1 = Path(tmp) / "out.docx.bak1"
            assert bak1.exists()

    def test_backup_zero_removes_all(self):
        """max_backups=0 should remove all numbered backups."""
        from docx_pipeline.config.defaults import get_template
        from docx_pipeline.config.schema import DocxPipelineConfig
        from docx_pipeline.converters.base import AbstractConverter

        data = get_template("default")
        data["backup"]["enabled"] = True
        data["backup"]["max_backups"] = 0
        cfg = DocxPipelineConfig.from_dict(data)

        class _TestConverter(AbstractConverter):
            def convert(self):
                from docx import Document
                return Document()

        with tempfile.TemporaryDirectory() as tmp:
            out = Path(tmp) / "out.docx"
            out.write_text("v1")
            # Pre-create some old backups
            for i in range(1, 6):
                (Path(tmp) / f"out.docx.bak{i}").write_text(f"old{i}")
            converter = _TestConverter(cfg)
            converter._rotate_backups(out)
            # All backups should be gone
            for i in range(1, 6):
                assert not (Path(tmp) / f"out.docx.bak{i}").exists()

    def test_backup_limit_reduction_cleans_excess(self):
        """Reducing max_backups from 5 to 2 should clean .bak3/4/5."""
        from docx_pipeline.config.defaults import get_template
        from docx_pipeline.config.schema import DocxPipelineConfig
        from docx_pipeline.converters.base import AbstractConverter

        data = get_template("default")
        data["backup"]["enabled"] = True
        data["backup"]["max_backups"] = 2
        cfg = DocxPipelineConfig.from_dict(data)

        class _TestConverter(AbstractConverter):
            def convert(self):
                from docx import Document
                return Document()

        with tempfile.TemporaryDirectory() as tmp:
            out = Path(tmp) / "out.docx"
            out.write_text("v1")
            # Pre-create 5 old backups
            for i in range(1, 6):
                (Path(tmp) / f"out.docx.bak{i}").write_text(f"old{i}")
            converter = _TestConverter(cfg)
            converter._rotate_backups(out)
            # Only .bak1 and .bak2 should survive
            assert (Path(tmp) / "out.docx.bak1").exists()
            assert (Path(tmp) / "out.docx.bak2").exists()
            for i in range(3, 6):
                assert not (Path(tmp) / f"out.docx.bak{i}").exists()

    def test_backup_non_consecutive_cleanup(self):
        """Non-consecutive backup numbers (.bak1, .bak3, .bak100) are all cleaned."""
        from docx_pipeline.config.defaults import get_template
        from docx_pipeline.config.schema import DocxPipelineConfig
        from docx_pipeline.converters.base import AbstractConverter

        data = get_template("default")
        data["backup"]["enabled"] = True
        data["backup"]["max_backups"] = 1
        cfg = DocxPipelineConfig.from_dict(data)

        class _TestConverter(AbstractConverter):
            def convert(self):
                from docx import Document
                return Document()

        with tempfile.TemporaryDirectory() as tmp:
            out = Path(tmp) / "out.docx"
            out.write_text("v1")
            # Non-consecutive: .bak1, .bak3, .bak100
            for n in (1, 3, 100):
                (Path(tmp) / f"out.docx.bak{n}").write_text(f"old{n}")
            converter = _TestConverter(cfg)
            converter._rotate_backups(out)
            # Only .bak1 should remain
            assert (Path(tmp) / "out.docx.bak1").exists()
            assert not (Path(tmp) / "out.docx.bak3").exists()
            assert not (Path(tmp) / "out.docx.bak100").exists()
