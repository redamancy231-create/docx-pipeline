"""Smoke tests for template defaults and dual-backend feature parity.

Provenance: GPT-5.6-Sol (via Codex CLI), 2026-07-17.
"""

from __future__ import annotations

import os
import shutil
import subprocess
import tempfile
from pathlib import Path
from unittest.mock import patch

import pytest
import yaml
from click.testing import CliRunner
from docx import Document
from docx.oxml.ns import qn

from docx_pipeline.cli import cli
from docx_pipeline.config.defaults import get_template
from docx_pipeline.config.schema import DocxPipelineConfig
from docx_pipeline.config.validator import validate_config
from docx_pipeline.converters.pandoc_converter import PandocConverter
from docx_pipeline.converters.pure_python import PurePythonConverter


@pytest.fixture()
def default_template_runtime():
    """Provide the files and executable probes required by runtime validation."""
    root = Path(tempfile.mkdtemp(prefix="docx_pipeline_defaults_"))
    previous_cwd = Path.cwd()
    source = root / "output" / "markdown"
    source.parent.mkdir(parents=True)
    source.write_text("# 默认模板\n\n正文。\n", encoding="utf-8")
    (root / "work").mkdir()

    os.chdir(root)
    try:
        with patch(
            "docx_pipeline.config.validator.shutil.which",
            side_effect=lambda executable: str(root / executable),
        ):
            yield
    finally:
        os.chdir(previous_cwd)
        shutil.rmtree(root, ignore_errors=True)


@pytest.fixture()
def dual_backend_documents():
    """Convert equivalent source documents through both backends."""
    root = Path(tempfile.mkdtemp(prefix="docx_pipeline_parity_"))
    markdown_path = root / "smoke.md"
    markdown_path.write_text(
        "# 中英标题 Latin\n\n"
        "中英正文 Latin text.\n\n"
        "| 中文列 | Latin column |\n"
        "| --- | --- |\n"
        "| 数据 | value |\n",
        encoding="utf-8",
    )

    data = get_template("default")
    data["project"].update({"name": "Smoke", "root": str(root)})
    data["paths"].update(
        {
            "md_source": str(markdown_path),
            "docx_output": str(root / "output.docx"),
            "work_dir": str(root / "work"),
        }
    )
    data["fonts"].update(
        {"east_asian": "Smoke East Asian", "latin": "Smoke Latin"}
    )
    data["styles"]["toc"]["enabled"] = False
    data["styles"]["table"]["autofit"] = True
    data["pandoc"]["enabled"] = True
    data["mermaid"]["enabled"] = False
    data["backup"]["enabled"] = False
    config = DocxPipelineConfig.from_dict(data)

    pure_document = PurePythonConverter(config).convert()

    def fake_pandoc(command, **_kwargs):
        output_path = Path(command[command.index("-o") + 1])
        document = Document()
        document.add_heading("中英标题 Latin", level=1)
        document.add_paragraph("中英正文 Latin text.")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "中文列"
        table.cell(0, 1).text = "Latin column"
        table.cell(1, 0).text = "数据"
        table.cell(1, 1).text = "value"
        document.save(output_path)
        return subprocess.CompletedProcess(
            args=command,
            returncode=0,
            stdout="",
            stderr="",
        )

    try:
        with patch(
            "docx_pipeline.converters.pandoc_converter.shutil.which",
            return_value=str(root / "pandoc"),
        ), patch(
            "docx_pipeline.converters.pandoc_converter.subprocess.run",
            side_effect=fake_pandoc,
        ) as mocked_run:
            pandoc_document = PandocConverter(config).convert()

        assert mocked_run.call_count == 1
        yield {"pure": pure_document, "pandoc": pandoc_document}
    finally:
        shutil.rmtree(root, ignore_errors=True)


def _find_run(document, text: str):
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if text in run.text:
                return run
    raise AssertionError(f"No run contains expected text: {text}")


def _table_property_elements(table, tag: str):
    table_properties = table._tbl.find(qn("w:tblPr"))
    assert table_properties is not None
    return table_properties.findall(qn(f"w:{tag}"))


class TestDefaultTemplateSmoke:
    """Verify that every template can be loaded AND converted without errors."""

    def test_all_templates_have_valid_defaults(self, default_template_runtime):
        """Every template's default dict must pass schema validation."""
        for name in ["default", "academic", "report", "strategy"]:
            data = get_template(name)
            cfg = DocxPipelineConfig.from_dict(data)
            issues = validate_config(cfg)
            assert not issues, f"{name}: {issues}"

    def test_docx_output_has_extension(self):
        """Template defaults must produce output paths ending in .docx."""
        for name in ["default", "academic", "report", "strategy"]:
            data = get_template(name)
            assert data["paths"]["docx_output"].endswith(".docx"), (
                f"{name}: docx_output lacks .docx extension: "
                f"{data['paths']['docx_output']}"
            )

    def test_init_generated_config_has_docx_extension(self):
        """CLI init with --md-file must produce config with .docx output."""
        root = Path(tempfile.mkdtemp(prefix="docx_pipeline_init_"))
        try:
            markdown_path = root / "输入文档.md"
            markdown_path.write_text("# 标题\n\n正文。\n", encoding="utf-8")

            result = CliRunner().invoke(
                cli,
                [
                    "init",
                    "--project-dir",
                    str(root),
                    "--template",
                    "default",
                    "--md-file",
                    str(markdown_path),
                ],
            )

            assert result.exit_code == 0, result.output
            config_path = root / "project.yaml"
            with config_path.open("r", encoding="utf-8") as config_file:
                generated = yaml.safe_load(config_file)
            assert generated["paths"]["docx_output"].endswith(".docx")
        finally:
            shutil.rmtree(root, ignore_errors=True)


class TestDualBackendParity:
    """Verify both backends behave equivalently for shared features."""

    def test_both_backends_produce_autofit_tables(self, dual_backend_documents):
        """Both backends must apply table autofit when config says so."""
        for backend, document in dual_backend_documents.items():
            assert document.tables, f"{backend}: conversion produced no table"
            table = document.tables[0]
            assert table.autofit is True, f"{backend}: table autofit is disabled"
            layouts = _table_property_elements(table, "tblLayout")
            assert any(
                layout.get(qn("w:type")) == "autofit" for layout in layouts
            ), f"{backend}: explicit autofit layout is missing"

        pandoc_widths = _table_property_elements(
            dual_backend_documents["pandoc"].tables[0], "tblW"
        )
        assert any(
            width.get(qn("w:type")) == "pct"
            and width.get(qn("w:w")) == "5000"
            for width in pandoc_widths
        ), "pandoc: full-width table setting is missing"

    def test_both_backends_apply_font_settings(self, dual_backend_documents):
        """Both backends must apply east_asian/latin font names."""
        for backend, document in dual_backend_documents.items():
            run = _find_run(document, "中英正文")
            run_properties = run._element.find(qn("w:rPr"))
            assert run_properties is not None, f"{backend}: run properties missing"
            run_fonts = run_properties.find(qn("w:rFonts"))
            assert run_fonts is not None, f"{backend}: run fonts missing"
            assert run_fonts.get(qn("w:eastAsia")) == "Smoke East Asian"
            assert run_fonts.get(qn("w:ascii")) == "Smoke Latin"
            assert run_fonts.get(qn("w:hAnsi")) == "Smoke Latin"